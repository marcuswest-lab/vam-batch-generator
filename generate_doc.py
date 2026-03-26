#!/usr/bin/env python3
"""
Generate a filled BAD Marketing .docx brief from Claude's output.

Usage:
  python3 generate_doc.py [paste_file]

If no file is given, reads from clipboard or prompts to paste.
Outputs a .docx file to the project directory.
"""

import sys
import os
import re
import shutil
import copy
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES = {
    'static': os.path.join(SCRIPT_DIR, '[Active] Static Copy Creative Brief _ Template.docx'),
    'video': os.path.join(SCRIPT_DIR, '[Active] Video Copy_Creative Brief _ Template.docx'),
    'copy': os.path.join(SCRIPT_DIR, '[Active] Body Copy Creative Brief _ Template.docx'),
}

# ---------------------------------------------------------------------------
# Value normalization: map common Claude outputs to valid dropdown values
# ---------------------------------------------------------------------------

LEAD_TYPE_ALIASES = {
    'proof': 'Proclamation',
    'pain': 'Problem-Solution',
    'curiosity': 'Secret',
    'problem': 'Problem-Solution',
    'solution': 'Problem-Solution',
    'problem solution': 'Problem-Solution',
    'story': 'Story',
}

VARIATION_TYPE_ALIASES = {
    'hook': 'Copy',          # "Hook" is not a valid static variation type → default to Copy
    'headline': 'Copy',
    'visual style': 'Visual',
    'visual': 'Visual',
}

STATUS_ALIASES = {
    'ready': 'Ready For Internal',
    'ready for review': 'Ready For Internal',
    'internal': 'Ready For Internal',
    'approved': 'Approved',
    'changes': 'Changes Required',
    'needs approval': 'Needs Client Approval',
    'client approval': 'Needs Client Approval',
}


def normalize_dropdown_value(value, dropdown_options, field_name=''):
    """Try to match value to a valid dropdown option, with alias fallback."""
    if not value:
        return value

    # Direct case-insensitive match
    for opt in dropdown_options:
        if value.strip().lower() == opt.lower():
            return opt

    # Try aliases based on field name
    val_lower = value.strip().lower()
    aliases = {}
    if 'lead type' in field_name.lower():
        aliases = LEAD_TYPE_ALIASES
    elif 'variation type' in field_name.lower():
        aliases = VARIATION_TYPE_ALIASES
    elif 'status' in field_name.lower():
        aliases = STATUS_ALIASES

    if val_lower in aliases:
        mapped = aliases[val_lower]
        # Verify the mapped value is actually in dropdown options
        for opt in dropdown_options:
            if mapped.lower() == opt.lower():
                return opt

    # Partial match (value is substring of option or vice versa)
    for opt in dropdown_options:
        if val_lower in opt.lower() or opt.lower() in val_lower:
            return opt

    # No match found — return original
    return value


# ---------------------------------------------------------------------------
# Parsing: extract structured data from Claude's text output
# ---------------------------------------------------------------------------

def detect_batch_type(text):
    """Detect batch type from output text."""
    text_lower = text.lower()
    if 'batch type: video' in text_lower or '=== video' in text_lower:
        return 'video'
    elif 'batch type: body' in text_lower or '=== copy' in text_lower or 'batch type: copy' in text_lower:
        return 'copy'
    else:
        return 'static'


def parse_field(text, field_name):
    """Extract a field value from 'Field Name: value' pattern.
    Uses [ \\t]* (horizontal whitespace) after the colon to avoid
    bridging newlines into the next field when the current field is empty."""
    patterns = [
        rf'(?:^|\n)\s*{re.escape(field_name)}:[ \t]*(.+?)(?:\n|$)',
        rf'(?:^|\n)\s*{re.escape(field_name)}[ \t]*:[ \t]*(.+?)(?:\n|$)',
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            val = m.group(1).strip()
            # Strip markdown bold markers
            val = re.sub(r'\*\*(.+?)\*\*', r'\1', val)
            return val
    return ''


def parse_multiline_field(text, field_name, stop_fields=None):
    """Extract a multi-line field value, stopping at the next known field or section marker."""
    if stop_fields is None:
        stop_fields = []

    # Build stop pattern from known fields
    stop_parts = [r'\n===', r'\n---\s*$']
    for sf in stop_fields:
        stop_parts.append(rf'\n\s*{re.escape(sf)}:')

    stop_pattern = '|'.join(stop_parts)

    pattern = rf'(?:^|\n)\s*{re.escape(field_name)}:\s*\n(.*?)(?={stop_pattern}|\Z)'
    m = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
    if m:
        return m.group(1).strip()

    # Try single-line fallback
    return parse_field(text, field_name)


def parse_overview(text):
    """Parse the OVERVIEW section."""
    overview = {}
    # Find the overview block
    ov_match = re.search(
        r'===\s*OVERVIEW\s*===\s*\n(.*?)(?====\s*(?:CREATIVE|VIDEO|COPY)\s+\d)',
        text, re.DOTALL | re.IGNORECASE
    )
    if not ov_match:
        # Try without === markers or with --- separator
        ov_match = re.search(
            r'OVERVIEW\s*\n(.*?)(?=(?:===\s*)?(?:CREATIVE|VIDEO|COPY)\s+\d)',
            text, re.DOTALL | re.IGNORECASE
        )
    if not ov_match:
        # Last resort: everything before first creative section
        ov_match = re.search(
            r'^(.*?)(?=(?:===\s*)?(?:CREATIVE|VIDEO|COPY)\s+1)',
            text, re.DOTALL | re.IGNORECASE
        )

    if ov_match:
        ov_text = ov_match.group(1)
        fields = [
            'AI Allowed?', 'Video Type', 'Photo Folder', 'Footage Folder', 'Reference',
            'Idea Name', 'Angle Name', 'Style Name', 'Task', 'General Notes',
            'Design Notes', 'Editing Notes', 'Link to Brand Assets', 'Any Other Relevant Assets',
            'Ratio Format', 'Ad Platform', 'Avatar', 'Brand Voice',
            'Net New/Iteration', 'Landing Page URL', 'Conversion Objective', 'Copywriter',
            'Copy Type'
        ]
        for field in fields:
            val = parse_field(ov_text, field)
            if val:
                overview[field] = val
        # Also try "Ratio Format(s)"
        val = parse_field(ov_text, 'Ratio Format(s)')
        if val:
            overview['Ratio Format(s)'] = val

    return overview


def parse_creatives(text, batch_type):
    """Parse individual creative sections."""
    creatives = []

    # Match section headers — support both === CREATIVE N === and --- separated blocks
    if batch_type == 'static':
        pattern = r'===\s*CREATIVE\s+(\d+)\s*===\s*\n(.*?)(?=\n===\s*(?:CREATIVE|BATCH)\s|$)'
    elif batch_type == 'video':
        pattern = r'===\s*VIDEO\s+(\d+)\s*===\s*\n(.*?)(?=\n===\s*(?:VIDEO|BATCH)\s|$)'
    else:  # copy
        pattern = r'===\s*COPY\s+(\d+)\s*===\s*\n(.*?)(?=\n===\s*(?:COPY|BATCH)\s|$)'

    matches = list(re.finditer(pattern, text, re.DOTALL | re.IGNORECASE))

    # Fallback: try without === markers (some Claude outputs use ### or **)
    if not matches:
        if batch_type == 'static':
            pattern = r'(?:#{1,3}\s*)?CREATIVE\s+(\d+)\s*(?:===)?\s*\n(.*?)(?=(?:#{1,3}\s*)?CREATIVE\s+\d|BATCH\s+SUMMARY|POST-GENERATION|$)'
        elif batch_type == 'video':
            pattern = r'(?:#{1,3}\s*)?VIDEO\s+(\d+)\s*(?:===)?\s*\n(.*?)(?=(?:#{1,3}\s*)?VIDEO\s+\d|BATCH\s+SUMMARY|POST-GENERATION|$)'
        else:
            pattern = r'(?:#{1,3}\s*)?COPY\s+(\d+)\s*(?:===)?\s*\n(.*?)(?=(?:#{1,3}\s*)?COPY\s+\d|BATCH\s+SUMMARY|POST-GENERATION|$)'
        matches = list(re.finditer(pattern, text, re.DOTALL | re.IGNORECASE))

    for m in matches:
        num = int(m.group(1))
        block = m.group(2)
        creative = {'number': num}

        if batch_type == 'static':
            simple_fields = ['File Name', 'File', 'Variation Type',
                             'Awareness Level', 'Lead Type', 'Status']
            multiline_fields = {
                'Notes': ['Design Notes', 'Variation Type', 'File', 'Awareness Level'],
                'Design Notes': ['Variation Type', 'Awareness Level', 'Lead Type'],
                'Copy': [],  # Last field — goes to end of block
            }
        elif batch_type == 'video':
            simple_fields = ['File Name', 'Video File', 'Variation Type',
                             'Awareness Level', 'Lead Type', 'Status']
            multiline_fields = {
                'Notes': ['Editing Notes', 'Variation Type'],
                'Editing Notes': ['Variation Type', 'Awareness Level'],
                'Lead Script': ['Body Script'],
                'Body Script': ['FACEBOOK', 'Status'],
            }
        else:  # copy
            simple_fields = ['Name', 'File Name', 'Variation Type',
                             'Awareness Level', 'Lead Type', 'Status', 'Headline']
            multiline_fields = {
                'Notes': ['Variation Type', 'Awareness Level'],
                'Body Copy': [],
            }

        # Parse simple single-line fields
        for field in simple_fields:
            val = parse_field(block, field)
            if val:
                creative[field] = val

        # Parse multi-line fields
        for field, stop_fields in multiline_fields.items():
            if field not in creative:
                val = parse_multiline_field(block, field, stop_fields)
                if val:
                    creative[field] = val

        # Fallback for Copy field: everything after "Copy:" to end of block
        if batch_type == 'static' and 'Copy' not in creative:
            copy_idx = block.lower().find('copy:')
            if copy_idx >= 0:
                creative['Copy'] = block[copy_idx + 5:].strip()

        # Fallback for Body Copy field
        if batch_type == 'copy' and 'Body Copy' not in creative:
            bc_idx = block.lower().find('body copy:')
            if bc_idx >= 0:
                creative['Body Copy'] = block[bc_idx + 10:].strip()

        # Use File Name as fallback for Name in copy batches
        if batch_type == 'copy' and 'Name' not in creative and 'File Name' in creative:
            creative['Name'] = creative['File Name']

        creatives.append(creative)

    return sorted(creatives, key=lambda c: c.get('number', 0))


# ---------------------------------------------------------------------------
# Auto-detect Variation Type from file names (isolate variables)
# ---------------------------------------------------------------------------

def auto_detect_variation_types(creatives, batch_type):
    """Auto-detect whether creatives are Copy or Visual variations.

    For static batches:
      - Extract visual style code from file names (part after '|')
      - If batch has more than one unique visual style → 'Visual' variations
      - If batch has only one visual style → 'Copy' variations

    This ensures proper variable isolation in batches.
    """
    if batch_type != 'static':
        return  # Only applies to static batches (Copy/Visual)

    # Extract visual style codes from file names
    visual_codes = []
    for c in creatives:
        fname = c.get('File Name', '')
        if '|' in fname:
            visual_code = fname.split('|', 1)[1].strip()
            visual_codes.append(visual_code)
        else:
            visual_codes.append('')

    unique_visuals = set(v for v in visual_codes if v)

    if len(unique_visuals) > 1:
        # Multiple visual styles → Visual variations
        print(f"  Auto-detect: {len(unique_visuals)} visual styles found ({', '.join(unique_visuals)}) → setting all to 'Visual'")
        for c in creatives:
            c['Variation Type'] = 'Visual'
    elif len(unique_visuals) == 1:
        # Single visual style → Copy variations
        print(f"  Auto-detect: 1 visual style ({list(unique_visuals)[0]}) → setting all to 'Copy'")
        for c in creatives:
            c['Variation Type'] = 'Copy'
    # If no visual codes found, leave whatever Claude set


# ---------------------------------------------------------------------------
# Fill document header (Client, Campaign Name, Date)
# ---------------------------------------------------------------------------

def fill_header(doc, claude_output, batch_type):
    """Replace [CLIENT], [CAMPAIGN/LAUNCH NAME], [DATE OF CREATION] in doc header."""
    batch_name = parse_field(claude_output, 'BATCH') or 'Batch'
    # Strip date suffixes from campaign name (e.g. "— Week of March 4, 2026")
    campaign_name = re.sub(r'\s*[—–-]\s*[Ww]eek\s+of\s+.*$', '', batch_name).strip()
    campaign_name = re.sub(r'\s*[—–-]\s*\d{1,2}[/-]\d{1,2}[/-]\d{2,4}.*$', '', campaign_name).strip()
    if not campaign_name:
        campaign_name = batch_name

    date_str = datetime.now().strftime('%m-%d-%Y')

    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            for run in para.runs:
                text = run.text
                if '[CLIENT]' in text or '[CAMPAIGN/LAUNCH NAME]' in text:
                    text = text.replace('[CLIENT]', 'Value Added Moving')
                    text = text.replace('[CAMPAIGN/LAUNCH NAME]', campaign_name)
                    run.text = text
                if '[DATE OF CREATION]' in text:
                    run.text = text.replace('[DATE OF CREATION]', date_str)


# ---------------------------------------------------------------------------
# .docx filling: use actual BAD Marketing template
# ---------------------------------------------------------------------------

def set_cell_text(cell, text):
    """Set cell text, preserving the first paragraph's formatting.
    Handles multi-line text by creating additional paragraphs.
    Clears ALL existing content (runs, SDTs, comments) to prevent leaks."""
    if not text:
        return

    lines = text.split('\n')

    # Get formatting from existing first paragraph
    first_para = cell.paragraphs[0]
    rPr_template = None
    pPr_template = None

    # Extract run properties from first available run (even inside SDTs)
    all_runs = first_para._element.findall('.//' + qn('w:r'))
    for r_elem in all_runs:
        rPr_elem = r_elem.find(qn('w:rPr'))
        if rPr_elem is not None:
            rPr_template = copy.deepcopy(rPr_elem)
            break

    pPr_elem = first_para._element.find(qn('w:pPr'))
    if pPr_elem is not None:
        pPr_template = copy.deepcopy(pPr_elem)

    # Clear the first paragraph completely — remove everything except pPr
    for child in list(first_para._element):
        if child.tag != qn('w:pPr'):
            first_para._element.remove(child)

    # Add new run with the first line of text
    run_elem = OxmlElement('w:r')
    if rPr_template is not None:
        run_elem.append(copy.deepcopy(rPr_template))
    t_elem = OxmlElement('w:t')
    t_elem.set(qn('xml:space'), 'preserve')
    t_elem.text = lines[0]
    run_elem.append(t_elem)
    first_para._element.append(run_elem)

    # Remove extra paragraphs (beyond the first)
    for i in range(len(cell.paragraphs) - 1, 0, -1):
        p_element = cell.paragraphs[i]._element
        p_element.getparent().remove(p_element)

    # Add remaining lines as new paragraphs
    for line in lines[1:]:
        new_p = OxmlElement('w:p')
        if pPr_template is not None:
            new_p.append(copy.deepcopy(pPr_template))
        run_elem = OxmlElement('w:r')
        if rPr_template is not None:
            run_elem.append(copy.deepcopy(rPr_template))
        t_elem = OxmlElement('w:t')
        t_elem.set(qn('xml:space'), 'preserve')
        t_elem.text = line
        run_elem.append(t_elem)
        new_p.append(run_elem)
        cell._tc.append(new_p)


def set_dropdown_value(cell, value, field_name=''):
    """Set a dropdown (structured document tag) value in a cell.
    Returns True if a dropdown was found and the value matched an option."""
    tc = cell._tc
    sdts = tc.findall('.//' + qn('w:sdt'))

    for sdt in sdts:
        sdtPr = sdt.find(qn('w:sdtPr'))
        if sdtPr is None:
            continue

        dropdown = sdtPr.find(qn('w:dropDownList'))
        comboBox = sdtPr.find(qn('w:comboBox'))
        dd = dropdown if dropdown is not None else comboBox
        if dd is None:
            continue

        # Collect all valid options
        options = []
        for item in dd.findall(qn('w:listItem')):
            display = item.get(qn('w:displayText')) or item.get(qn('w:value'))
            if display:
                options.append(display)

        # Normalize the value to match a valid option
        normalized = normalize_dropdown_value(value, options, field_name)

        # Find the matching option
        matched_display = None
        matched_value = None
        for item in dd.findall(qn('w:listItem')):
            display = item.get(qn('w:displayText')) or item.get(qn('w:value'))
            item_val = item.get(qn('w:value'))
            if display and normalized.lower() == display.lower():
                matched_display = display
                matched_value = item_val or display
                break

        if matched_display:
            # Update lastValue attribute on the dropdown element
            dd.set(qn('w:lastValue'), matched_value)

            # Update the display text in the SDT content
            sdtContent = sdt.find(qn('w:sdtContent'))
            if sdtContent is not None:
                # Try runs directly under sdtContent (common structure)
                runs = sdtContent.findall(qn('w:r'))
                if runs:
                    for r in runs:
                        for t in r.findall(qn('w:t')):
                            t.text = matched_display
                    return True

                # Try runs inside paragraphs under sdtContent
                for p in sdtContent.findall(qn('w:p')):
                    for r in p.findall(qn('w:r')):
                        for t in r.findall(qn('w:t')):
                            t.text = matched_display
                    return True

                # Last resort: create a new run with the text
                run_elem = OxmlElement('w:r')
                t_elem = OxmlElement('w:t')
                t_elem.set(qn('xml:space'), 'preserve')
                t_elem.text = matched_display
                run_elem.append(t_elem)
                sdtContent.append(run_elem)
                return True

        # Value doesn't match any option — log warning
        print(f"  WARNING: '{value}' (normalized: '{normalized}') not in dropdown options {options} for {field_name}")

    # No SDT found — set as plain text
    set_cell_text(cell, value)
    return False


# ---------------------------------------------------------------------------
# Template filling functions — one per batch type
# ---------------------------------------------------------------------------

def fill_static_doc(doc, overview, creatives):
    """Fill static brief template.

    STATIC OVERVIEW TABLE (Table 0, 19 rows):
      R0:  OVERVIEW (header)
      R1:  AI Allowed?         [dropdown]
      R2:  Photo Folder
      R3:  Reference
      R4:  Idea Name
      R5:  Angle Name
      R6:  Style Name
      R7:  Task
      R8:  General Notes
      R9:  Design Notes
      R10: Link to Brand Assets
      R11: Ratio Format(s)
      R12: Ad Platform          [dropdown]
      R13: Avatar
      R14: Brand Voice
      R15: Net New/Iteration    [dropdown]
      R16: Landing Page URL
      R17: Conversion Objective
      R18: Copywriter

    STATIC CREATIVE TABLES (Tables 1-5, 10 rows each):
      R0:  # (header)
      R1:  File Name
      R2:  File
      R3:  Notes
      R4:  Design Notes
      R5:  Variation Type       [dropdown: Copy, Visual]
      R6:  Awareness Level      [dropdown]
      R7:  Lead Type            [dropdown]
      R8:  Status               [dropdown]
      R9:  Copy
    """
    table = doc.tables[0]

    overview_map = {
        1:  ('AI Allowed?', True),
        2:  ('Photo Folder', False),
        3:  ('Reference', False),
        4:  ('Idea Name', False),
        5:  ('Angle Name', False),
        6:  ('Style Name', False),
        7:  ('Task', False),
        8:  ('General Notes', False),
        9:  ('Design Notes', False),
        10: ('Link to Brand Assets', False),
        11: ('Ratio Format(s)', False),
        12: ('Ad Platform', True),
        13: ('Avatar', False),
        14: ('Brand Voice', False),
        15: ('Net New/Iteration', True),
        16: ('Landing Page URL', False),
        17: ('Conversion Objective', False),
        18: ('Copywriter', False),
    }

    _fill_overview(table, overview, overview_map)

    creative_field_map = {
        1: ('File Name', False),
        2: ('File', False),
        3: ('Notes', False),
        4: ('Design Notes', False),
        5: ('Variation Type', True),
        6: ('Awareness Level', True),
        7: ('Lead Type', True),
        8: ('Status', True),
        9: ('Copy', False),
    }

    _fill_creatives(doc, creatives, creative_field_map)


def fill_video_doc(doc, overview, creatives):
    """Fill video brief template.

    VIDEO OVERVIEW TABLE (Table 0, 20 rows):
      R0:  OVERVIEW (header)
      R1:  Video Type            [dropdown]
      R2:  AI Allowed?           [dropdown]
      R3:  Footage Folder
      R4:  Idea Name
      R5:  Angle Name
      R6:  Style Name
      R7:  Task
      R8:  General Notes
      R9:  Editing Notes
      R10: Link to Brand Assets
      R11: Any Other Relevant Assets
      R12: Ratio Format(s)
      R13: Ad Platform           [dropdown]
      R14: Avatar
      R15: Brand Voice
      R16: Net New/Iteration     [dropdown]
      R17: Landing Page URL
      R18: Conversion Objective
      R19: Copywriter

    VIDEO CREATIVE TABLES (Tables 1-5, 11 rows each):
      R0:  # (header)
      R1:  File Name
      R2:  Video File
      R3:  Notes
      R4:  Editing Notes
      R5:  Variation Type        [dropdown: Lead, Pattern Interrupt, Body, CTA]
      R6:  Awareness Level       [dropdown]
      R7:  Lead Type             [dropdown]
      R8:  Status                [dropdown]
      R9:  Lead Script
      R10: Body Script
    """
    table = doc.tables[0]

    overview_map = {
        1:  ('Video Type', True),
        2:  ('AI Allowed?', True),
        3:  ('Footage Folder', False),
        4:  ('Idea Name', False),
        5:  ('Angle Name', False),
        6:  ('Style Name', False),
        7:  ('Task', False),
        8:  ('General Notes', False),
        9:  ('Editing Notes', False),
        10: ('Link to Brand Assets', False),
        11: ('Any Other Relevant Assets', False),
        12: ('Ratio Format(s)', False),
        13: ('Ad Platform', True),
        14: ('Avatar', False),
        15: ('Brand Voice', False),
        16: ('Net New/Iteration', True),
        17: ('Landing Page URL', False),
        18: ('Conversion Objective', False),
        19: ('Copywriter', False),
    }

    _fill_overview(table, overview, overview_map)

    creative_field_map = {
        1:  ('File Name', False),
        2:  ('Video File', False),
        3:  ('Notes', False),
        4:  ('Editing Notes', False),
        5:  ('Variation Type', True),
        6:  ('Awareness Level', True),
        7:  ('Lead Type', True),
        8:  ('Status', True),
        9:  ('Lead Script', False),
        10: ('Body Script', False),
    }

    _fill_creatives(doc, creatives, creative_field_map)


def fill_copy_doc(doc, overview, creatives):
    """Fill body copy brief template.

    COPY OVERVIEW TABLE (Table 0, 12 rows):
      R0:  OVERVIEW (header)
      R1:  AI Allowed?           [dropdown]
      R2:  Idea Name
      R3:  Angle Name
      R4:  Copy Type             [dropdown]
      R5:  Task
      R6:  General Notes
      R7:  Ad Platform           [dropdown]
      R8:  Net New/Iteration     [dropdown]
      R9:  Landing Page URL
      R10: Conversion Objective
      R11: Copywriter

    COPY CREATIVE TABLES (Tables 1-5, 9 rows each):
      R0:  # (header)
      R1:  Name
      R2:  Notes
      R3:  Variation Type        [dropdown: Lead, Body, CTA, Full]
      R4:  Awareness Level       [dropdown]
      R5:  Lead Type             [dropdown]
      R6:  Status                [dropdown]
      R7:  Headline
      R8:  Body Copy
    """
    table = doc.tables[0]

    overview_map = {
        1:  ('AI Allowed?', True),
        2:  ('Idea Name', False),
        3:  ('Angle Name', False),
        4:  ('Copy Type', True),
        5:  ('Task', False),
        6:  ('General Notes', False),
        7:  ('Ad Platform', True),
        8:  ('Net New/Iteration', True),
        9:  ('Landing Page URL', False),
        10: ('Conversion Objective', False),
        11: ('Copywriter', False),
    }

    _fill_overview(table, overview, overview_map)

    creative_field_map = {
        1: ('Name', False),
        2: ('Notes', False),
        3: ('Variation Type', True),
        4: ('Awareness Level', True),
        5: ('Lead Type', True),
        6: ('Status', True),
        7: ('Headline', False),
        8: ('Body Copy', False),
    }

    _fill_creatives(doc, creatives, creative_field_map, name_fallback=True)


# ---------------------------------------------------------------------------
# Shared helpers for filling
# ---------------------------------------------------------------------------

def _fill_overview(table, overview, overview_map):
    """Fill overview table using the field→row mapping."""
    for row_idx, (field, is_dropdown) in overview_map.items():
        if row_idx >= len(table.rows):
            print(f"  WARNING: Row {row_idx} out of range (table has {len(table.rows)} rows)")
            continue

        val = overview.get(field, '')

        # Try alternate field names
        if not val and field == 'Ratio Format(s)':
            val = overview.get('Ratio Format', '')
        if not val and field == 'Footage Folder':
            val = overview.get('Photo Folder', '')

        if val:
            cell = table.rows[row_idx].cells[1]
            if is_dropdown:
                set_dropdown_value(cell, val, field_name=field)
            else:
                set_cell_text(cell, val)


def _fill_creatives(doc, creatives, field_map, name_fallback=False):
    """Fill creative tables 1-5."""
    for i, creative in enumerate(creatives[:5]):
        table_idx = i + 1
        if table_idx >= len(doc.tables):
            print(f"  WARNING: Table {table_idx} not found (doc has {len(doc.tables)} tables)")
            break

        ctable = doc.tables[table_idx]
        print(f"  Filling Creative {creative.get('number', i+1)}...")

        for row_idx, (field, is_dropdown) in field_map.items():
            if row_idx >= len(ctable.rows):
                print(f"    WARNING: Row {row_idx} out of range (table has {len(ctable.rows)} rows)")
                break

            val = creative.get(field, '')

            # Name fallback for copy batches
            if not val and name_fallback and field == 'Name':
                val = creative.get('File Name', '')

            if val:
                cell = ctable.rows[row_idx].cells[1]
                if is_dropdown:
                    set_dropdown_value(cell, val, field_name=field)
                else:
                    set_cell_text(cell, val)


# ---------------------------------------------------------------------------
# Main generation pipeline
# ---------------------------------------------------------------------------

def generate_doc(claude_output, output_path=None):
    """Main function: parse Claude output and generate filled .docx."""

    # Detect batch type
    batch_type = detect_batch_type(claude_output)
    print(f"Detected batch type: {batch_type}")

    # Parse
    overview = parse_overview(claude_output)
    creatives = parse_creatives(claude_output, batch_type)
    print(f"Parsed overview with {len(overview)} fields:")
    for k, v in overview.items():
        print(f"  {k}: {v[:80]}{'...' if len(v) > 80 else ''}")
    print(f"Parsed {len(creatives)} creatives")
    for c in creatives:
        print(f"  Creative {c.get('number')}: {list(c.keys())}")

    if not creatives:
        print("\nERROR: No creatives found in the output.")
        print("Make sure the output contains === CREATIVE/VIDEO/COPY N === sections.")
        sys.exit(1)

    # Copy template
    template_path = TEMPLATES[batch_type]
    if not os.path.exists(template_path):
        print(f"ERROR: Template not found: {template_path}")
        sys.exit(1)

    if not output_path:
        # Extract batch name for filename
        batch_name = parse_field(claude_output, 'BATCH')
        if not batch_name:
            batch_name = f"VAM_{batch_type.title()}_Batch"
        # Clean filename
        safe_name = re.sub(r'[^\w\s-]', '', batch_name).strip()[:60]
        safe_name = re.sub(r'\s+', '_', safe_name)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M')
        output_path = os.path.join(SCRIPT_DIR, f"{safe_name}_{timestamp}.docx")

    shutil.copy2(template_path, output_path)

    # Auto-detect variation types (must happen before filling)
    print("\nDetecting variation types...")
    auto_detect_variation_types(creatives, batch_type)

    # Fill
    doc = Document(output_path)

    # Remove cover page / title paragraph (first paragraph before the overview table)
    if doc.paragraphs:
        first_p = doc.paragraphs[0]._element
        parent = first_p.getparent()
        # Also remove any trailing page break immediately after
        next_sib = first_p.getnext()
        parent.remove(first_p)
        # If the next element is an empty paragraph (page break spacer), remove it too
        if next_sib is not None and next_sib.tag == qn('w:p'):
            text_content = ''.join(t.text or '' for t in next_sib.findall('.//' + qn('w:t')))
            if not text_content.strip():
                parent.remove(next_sib)

    # Fill header fields (Client, Campaign Name, Date)
    print("Filling header fields...")
    fill_header(doc, claude_output, batch_type)

    print(f"\nFilling {batch_type} template...")
    if batch_type == 'static':
        fill_static_doc(doc, overview, creatives)
    elif batch_type == 'video':
        fill_video_doc(doc, overview, creatives)
    else:
        fill_copy_doc(doc, overview, creatives)

    doc.save(output_path)
    print(f"\n✓ Generated: {output_path}")
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    if len(sys.argv) > 1:
        # Read from file
        filepath = sys.argv[1]
        if not os.path.exists(filepath):
            print(f"ERROR: File not found: {filepath}")
            sys.exit(1)
        with open(filepath, 'r') as f:
            text = f.read()
        print(f"Read {len(text)} chars from {filepath}")
    else:
        # Try clipboard first (macOS)
        try:
            import subprocess
            result = subprocess.run(['pbpaste'], capture_output=True, text=True)
            if result.returncode == 0 and len(result.stdout) > 100:
                text = result.stdout
                print(f"Read {len(text)} chars from clipboard")
            else:
                raise Exception("Clipboard empty or too short")
        except Exception:
            print("Paste Claude's output below (press Ctrl+D when done):")
            text = sys.stdin.read()

    generate_doc(text)
